from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PIL import Image
import os

def cal(image1):
    img=Image.open(image1)
    width, height = img.size
    iw= 0.0104166667 * width
    ih=0.0104166667 * height

    pw=5.90551
    ph=8.66142

    pr=ph/pw
    ir=ih/iw

    dw=iw-pw
    dh=ir*dw

    mw=pw
    mh=ih-dh

    if mw>pw or mh>ph:
        dh=ih-ph
        dw=dh/ir
        mh=ph
        mw=iw-dw
    return mw,mh

def align():
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

document = Document()

name = document.add_paragraph().add_run("Your Images!")
name.font.size = Pt(28)
name.bold = True
name.underline = True
name.font.name = "Algerian"
align()

a=os.listdir()

notlst=[]

for i in a:
    try:
        mw,mh=cal(i)
        document.add_picture(i,width=Inches(mw), height=Inches(mh))
        align()
    except:
        notlst.append(i)

if notlst==[]:
    print("Your Work is Done! :-)")
else:
    print("Your Work is Done! But Due to Some reasons couldn't add the Files mentioned below! ")
    print(notlst)

document.save("images.docx")
